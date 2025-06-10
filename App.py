from flask import Flask, render_template, flash, request, session, send_file, jsonify
from flask import render_template, redirect, url_for, request
import mysql.connector
from docx import Document
from ecies.utils import generate_key
from ecies import encrypt, decrypt
import base64, os, sys

import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from collections import Counter
import string
import docx  # for reading .docx files
import fitz  # PyMuPDF, for reading .pdf files

# Ensure necessary NLTK packages are downloaded
nltk.download('punkt')
nltk.download('stopwords')

app = Flask(__name__)
app.config['DEBUG']
app.config['SECRET_KEY'] = '7d441f27d441f27567d441f2b6176a'


@app.route("/")
def homepage():
    return render_template('index.html')


@app.route("/ServerLogin")
def ServerLogin():
    return render_template('ServerLogin.html')


@app.route('/NewUser')
def NewUser():
    return render_template('NewUser.html')


@app.route('/UserLogin')
def UserLogin():
    return render_template('UserLogin.html')


@app.route("/serverlogin", methods=['GET', 'POST'])
def serverlogin():
    if request.method == 'POST':
        if request.form['uname'] == 'admin' and request.form['password'] == 'admin':

            conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
            cur = conn.cursor()
            cur.execute("SELECT * FROM regtb ")
            data = cur.fetchall()


            return render_template('ServerHome.html', data=data)

        else:
            flash('Username or Password is wrong')
            return render_template('ServerLogin.html')


@app.route("/ServerHome")
def ServerHome():
    conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
    cur = conn.cursor()
    cur.execute("SELECT * FROM regtb ")
    data = cur.fetchall()
    return render_template('ServerHome.html', data=data)

@app.route("/newuser", methods=['GET', 'POST'])
def newuser():
    if request.method == 'POST':
        uname = request.form['uname']
        mobile = request.form['mobile']
        email = request.form['email']
        address = request.form['address']
        username = request.form['username']
        password = request.form['password']

        conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
        cursor = conn.cursor()
        cursor.execute("SELECT * from regtb where username='" + username + "'  ")
        data = cursor.fetchone()
        if data is None:
            conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO regtb VALUES ('','" + uname + "','" + mobile + "','" + email + "','" + address + "','" +
                username + "','" + password + "')")
            conn.commit()
            conn.close()

            flash('Record Saved!')
            return render_template('NewUser.html')
        else:
            flash('Already Register This  UserName!')
            return render_template('NewUser.html')


@app.route("/userlogin", methods=['GET', 'POST'])
def userlogin():
    if request.method == 'POST':

        username = request.form['uname']
        password = request.form['password']

        session['uname'] = request.form['uname']

        conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
        cursor = conn.cursor()
        cursor.execute("SELECT * from regtb where username='" + username + "' and Password='" + password + "' ")
        data = cursor.fetchone()
        if data is None:

            flash("UserName or Password Incorrect..!")
            return render_template('UserLogin.html')

        else:
            conn = mysql.connector.connect(user='root', password='', host='localhost',
                                           database='1resumebuliderdb')
            cur = conn.cursor()
            cur.execute("SELECT * FROM regtb where username='" + session['uname'] + "'")
            data1 = cur.fetchall()
            flash('Login Successfully')
            return render_template('UserHome.html', data=data1)





@app.route("/Chat")
def Chat():
    return render_template('Chat.html')


import google.generativeai as genai

genai.configure(api_key='AIzaSyClK2fN-1M0eY-zcKQqXsfX85RVBYGSFfo')


def get_completion(out):
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(out)
    print(response.text)
    return response.text


@app.route("/get1")
def get_bot_response():
    userText = request.args.get('msg')
    print(userText)
    response = get_completion(userText)

    # return str(bot.get_response(userText))
    return response


@app.route("/UserHome")
def UserHome():
    conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
    cur = conn.cursor()
    cur.execute("SELECT * FROM regtb where UserName='" + session['uname'] + "'")
    data = cur.fetchall()
    return render_template('UserHome.html', data=data)


@app.route("/UserFileUpload")
def UserFileUpload():
    return render_template('UserFileUpload.html', uname=session['uname'])


import re

def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text


def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)

    return "\n".join([para.text for para in doc.paragraphs])

def extract_details(text):
    print(text)
    name = text.split('\n')[0].strip()  # Assuming first line is name

    email = re.search(r'[\w\.-]+@[\w\.-]+', text)
    phone = re.search(r'\+?\d[\d\s-]{8,15}', text)
    address = re.findall(r'\b(karur|salem|namakkal|erode|chennai|trichy|dindugal|madurai)\b', text,re.IGNORECASE)
    skills = re.findall(r'\b(Python|Java|C\+\+|HTML|CSS|JavaScript|SQL|Photoshop|Figma|C Programming)\b', text, re.IGNORECASE)
    soft_skills = re.findall(r'\b(Communication|Teamwork|Emotional Intelligence|Leadership|AdaptabilityTime Management|Problem-Solving)\b', text,
                        re.IGNORECASE)

    education = []
    edu_found = False
    edu_lines = text.split('\n')
    temp_entry = []

    for line in edu_lines:
        stripped = line.strip()

        if not edu_found and any(kw in stripped for kw in
                                 ['Education', 'B.E', 'B.Tech', 'HSC', 'SSLC', 'Higher Secondary', 'Secondary School',
                                  'Degree']):
            edu_found = True
            continue

        if edu_found:
            if stripped == "" or "skill" in stripped.lower():
                break

            temp_entry.append(stripped)

            if len(temp_entry) == 2:
                # Line 1: College name with bullet
                # Line 2: Degree + Duration, indented
                formatted = f"• {temp_entry[0]}\n  {temp_entry[1]}"
                education.append(formatted)
                temp_entry = []

            if len(education) == 3:
                break

    education_str = "\n".join(education)

    lines = text.split('\n')
    objective = ""
    for i, line in enumerate(lines):
        if any(kw in line.lower() for kw in ['objective', 'summary', 'career objective']):
            # Start grabbing the next lines
            j = i + 1
            while j < len(lines) and not any(
                    kw in lines[j].lower() for kw in ['education', 'degree', 'school', 'university']):
                objective += " " + lines[j].strip()
                if objective.count('.') >= 2:
                    break
                j += 1
            objective = objective.strip()
            break

        if any(kw in line.lower() for kw in ['education', 'degree', 'school', 'university']):
            break

    # Clean Experience section - no symbols or bullets
    experience_section = ""
    experience_started = False
    experience_lines = []

    for i, line in enumerate(lines):
        lower_line = line.lower().strip()

        if not experience_started and any(word in lower_line for word in ['experience', 'internship', 'worked']):
            experience_started = True
            continue  # Skip the header line

        if experience_started:
            # Stop at the next section
            if any(kw in lower_line for kw in ['education', 'skill', 'project', 'objective']):
                break

            # Clean unwanted bullet symbols like •, ●, ▪,  etc.
            clean_line = re.sub(r'^[^a-zA-Z0-9]*', '', line).strip()

            if clean_line:
                experience_lines.append(clean_line)

    experience_section = "\n".join(experience_lines)

    # Clean Experience section - no symbols or bullets
    certification_section = ""
    certification_started = False
    certification_lines = []

    for i, line in enumerate(lines):
        lower_line = line.lower().strip()

        if not certification_started and any(
                word in lower_line for word in ['certification', 'certifications', 'certificate', 'certificate','certificates']):
            certification_started = True
            continue  # Skip the header line

        if certification_started:
            # Stop at the next section
            if any(kw in lower_line for kw in ['education', 'skill', 'project', 'objective','honors','awards']):
                break

            # Clean unwanted bullet symbols like •, ●, ▪, • etc.
            clean_line = re.sub(r'^[^a-zA-Z0-9]*', '', line).strip()

            if clean_line:
                certification_lines.append(clean_line)

    certification_section = "\n".join(certification_lines)

    def extract_experience(doc_path):
        doc = Document(doc_path)
        full_text = []

        # Read all text from resume
        for para in doc.paragraphs:
            full_text.append(para.text.strip())

        resume_text = "\n".join(full_text)

        if 'Experience' in resume_text:
            experience_text = resume_text.split('Experience', 1)[1]
        elif 'Internship' in resume_text:
            experience_text = resume_text.split('Experience', 1)[1]
        else:
            return "No Experience section found."

        # Stop reading after next section like Projects, Certifications, Skills
        for section in ['Projects', 'Certifications', 'Skills', 'Education', 'Honors', 'Activities']:
            if section in experience_text:
                experience_text = experience_text.split(section)[0]

        return experience_text.strip()



    return {
        'name': name,
        'email': email.group() if email else '',
        'phone': phone.group() if phone else '',
        'address': ', '.join(set(address)),
        'skills': ', '.join(set(skills)),
        'soft_skills': ', '.join(set(soft_skills)),
        'education': education_str.strip(),
        'experience': experience_section.strip(),
        'objective': objective.strip(),
        'certification': certification_section.strip(),
    }

@app.route("/ufileupload", methods=['GET', 'POST'])
def ufileupload():
    if request.method == 'POST':
        uname = session['uname']

        file = request.files['file']
        import random
        fnew = random.randint(111, 999)
        savename = str(fnew) + file.filename

        file.save("static/upload/" + savename)

        file_path="static/upload/" + savename

        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(file_path)
        else:
            return "Unsupported file format"

        details = extract_details(text)
        print(details)
        return render_template('simple_resume.html',details=details)


@app.route('/UFileInfo')
def UFileInfo():
    conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
    cur = conn.cursor()
    cur.execute("SELECT * FROM filetb where username='" + session['uname'] + "'")
    data1 = cur.fetchall()
    return render_template('UFileInfo.html', data=data1)


@app.route('/USearch')
def USearch():
    return render_template('USearch.html')


@app.route("/search", methods=['GET', 'POST'])
def search():
    if request.method == 'POST':
        fname = request.form['uname']
        uname = session['uname']
        conn = mysql.connector.connect(user='root', password='', host='localhost',
                                       database='1resumebuliderdb')
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM filetb where (FileName like '%" + fname + "%' or FileInfo like '%" + fname + "%' or keyword like '%" + fname + "%' ) and UserName='" + uname + "'")
        data1 = cur.fetchall()
        # flash('Login Successfully')
        return render_template('USearch.html', data=data1)


@app.route("/search1", methods=['GET', 'POST'])
def search1():
    if request.method == 'POST':
        fname = request.form['uname']

        conn = mysql.connector.connect(user='root', password='', host='localhost',
                                       database='1resumebuliderdb')
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM filetb where (FileName like '%" + fname + "%' or FileInfo like '%" + fname + "%') ")
        data1 = cur.fetchall()
        # flash('Login Successfully')
        return render_template('UserHome1.html', data=data1)


@app.route("/Download1")
def Download1():
    lid = request.args.get('lid')

    conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
    cursor = conn.cursor()
    cursor.execute("SELECT  *  FROM  filetb where  id='" + str(lid) + "'")
    data = cursor.fetchone()
    if data:
        fname = data[3]
        uname = data[1]

        conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
        cursor = conn.cursor()
        cursor.execute("SELECT * from regtb where username='" + uname + "'  ")
        data1 = cursor.fetchone()
        if data1:
            sendmail(data1[3], "Unknown User Access Your File..!")

        newfilepath1 = './static/dupload/' + str(fname)

        return send_file(newfilepath1, as_attachment=True)

@app.route("/Download")
def Download():
    lid = request.args.get('lid')

    conn = mysql.connector.connect(user='root', password='', host='localhost', database='1resumebuliderdb')
    cursor = conn.cursor()
    cursor.execute("SELECT  *  FROM  filetb where  id='" + str(lid) + "'")
    data = cursor.fetchone()
    if data:
        fname = data[3]
        privhex = data[6]
        #newfilepath1 = './static/upload/' + str(fname)

        filepath = "./static/Encrypt/" + fname
        head, tail = os.path.split(filepath)

        newfilepath1 = './static/Encrypt/' + str(tail)
        newfilepath2 = './static/Decrypt/' + str(tail)

        data = 0
        with open(newfilepath1, "rb") as File:
            data = base64.b64decode(File.read())

        print(data)
        decrypted_secp = decrypt(privhex, data)
        print("\nDecrypted:", decrypted_secp)
        with open(newfilepath2, "wb") as DFile:
            DFile.write(base64.b64decode(decrypted_secp))

        return send_file(newfilepath2, as_attachment=True)



        #return send_file(newfilepath1, as_attachment=True)


def sendmail(Mailid, message):
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    fromaddr = "projectmailm@gmail.com"
    toaddr = Mailid

    # instance of MIMEMultipart
    msg = MIMEMultipart()

    # storing the senders email address
    msg['From'] = fromaddr

    # storing the receivers email address
    msg['To'] = toaddr

    # storing the subject
    msg['Subject'] = "Alert"

    # string to store the body of the mail
    body = message

    # attach the body with the msg instance
    msg.attach(MIMEText(body, 'plain'))

    # creates SMTP session
    s = smtplib.SMTP('smtp.gmail.com', 587)

    # start TLS for security
    s.starttls()

    # Authentication
    s.login(fromaddr, "qmgn xecl bkqv musr")

    # Converts the Multipart msg into a string
    text = msg.as_string()

    # sending the mail
    s.sendmail(fromaddr, toaddr, text)

    # terminating the session
    s.quit()

@app.route('/admin_search', methods=['POST'])
def admin_search():
    try:
        import os
        from PyPDF2 import PdfReader
        import re

        data = request.get_json(force=True)
        search_phrase = data.get('domain', '').strip().lower()

        if not search_phrase:
            return jsonify({'result': 'Please enter a skill or domain.'})

        def clean_text(txt):
            txt = txt.lower()
            txt = re.sub(r'[\n\r]', ' ', txt)  # replace newlines
            txt = re.sub(r'[^a-z0-9\s]', ' ', txt)  # remove punctuation
            txt = re.sub(r'\s+', ' ', txt)  # normalize spacing
            return txt.strip()

        matched_candidates = []
        UPLOAD_FOLDER = 'static/upload/'

        for filename in os.listdir(UPLOAD_FOLDER):
            if filename.lower().endswith('.pdf'):
                path = os.path.join(UPLOAD_FOLDER, filename)
                try:
                    with open(path, "rb") as f:
                        reader = PdfReader(f)
                        raw_text = ''
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                raw_text += page_text
                        cleaned_resume = clean_text(raw_text)
                        cleaned_search = clean_text(search_phrase)
                        if cleaned_search in cleaned_resume:
                            matched_candidates.append(filename.rsplit('.', 1)[0].title())
                except Exception as e:
                    print(f"Error reading file: {filename}, {e}")

        if matched_candidates:
            return jsonify({'result': "Experienced persons found:\n" + "\n".join(matched_candidates)})
        else:
            return jsonify({'result': f"Sorry, no persons found for domain: {search_phrase}"})
    except Exception as e:
        print("ERROR in /admin_search:", e)
        return jsonify({'result': 'Internal Server Error'}), 500

if __name__ == '__main__':
    app.run(debug=True, use_reloader=True)
